# ✍️ Motley Fool AI Copywriter — Pro Edition v4.1
# ----------------------------------------------------------
# • Built‑in “Internal Plan” stage
# • JSON {plan, copy} separation
# • Dynamic word‑count control tied to length dropdown
# • Auto‑expand via self‑QA when draft is too short
# ----------------------------------------------------------

import time, json
from io import BytesIO
from textwrap import dedent

import streamlit as st
from openai import OpenAI
from docx import Document

# ────────────────────────────────────────────────────────────
# 0.  Global toggles
# ────────────────────────────────────────────────────────────
USE_STREAMING = False     # live token stream
AUTO_QA      = True       # self‑critique & auto‑fix loop

# ---- Model & token ceiling ---------------------------------
MAX_OUTPUT_TOKENS = 10_000   # safe for GPT‑4‑32k context

# ---- Length buckets (words) --------------------------------
LENGTH_RULES = {
    "📏 Short (100–200 words)":        (100, 220),
    "📐 Medium (200–500 words)":       (200, 550),
    "📖 Long (500–1500 words)":        (500, 1600),
    "📚 Extra Long (1500–3000 words)": (1500, 3200),
    "📜 Scrolling Monster (3000+ words)": (3000, None),  # None = open‑ended
}

# ────────────────────────────────────────────────────────────
# 1.  OpenAI client & config
# ────────────────────────────────────────────────────────────
client = OpenAI(api_key=st.secrets.openai_api_key)
OPENAI_MODEL = st.secrets.get("openai_model", "gpt-4.1")

# ────────────────────────────────────────────────────────────
# 2.  Streamlit page & CSS
# ────────────────────────────────────────────────────────────
st.set_page_config(page_title="✍️ Motley Fool AI Copywriter",
                   initial_sidebar_state="expanded")
st.title("✍️ Motley Fool AI Copywriter")

st.markdown("""
<style>
div.stButton>button { width:100%; }
h2, h3   { margin-top:1.1em; }
ul       { margin-left:1.3em; }
strong   { color:#CF7F00; }
</style>
""", unsafe_allow_html=True)

# ────────────────────────────────────────────────────────────
# 3.  Session helpers
# ────────────────────────────────────────────────────────────
def _init(**defaults):
    for k, v in defaults.items():
        st.session_state.setdefault(k, v)

_init(generated_copy="", adapted_copy="", internal_plan="", length_choice="")

def line(label: str, value: str) -> str:
    return f"- {label}: {value}\n" if value.strip() else ""

# ────────────────────────────────────────────────────────────
# 4.  Prompt components
# ────────────────────────────────────────────────────────────
COUNTRY_RULES = {
    "Australia":      "Use Australian English, prices in AUD, reference the ASX.",
    "United Kingdom": "Use British English, prices in GBP, reference the FTSE.",
    "Canada":         "Use Canadian English, prices in CAD, reference the TSX.",
    "United States":  "Use American English, prices in USD, reference the S&P 500.",
}

SYSTEM_PROMPT = dedent("""
You are The Motley Fool’s senior direct‑response copy chief.

• Voice: plain English, optimistic, inclusive, lightly playful but always expert.
• Draw from Ogilvy clarity, Sugarman narrative, Halbert urgency, Cialdini persuasion.
• Use **Markdown headings** (##, ###) and standard `-` bullets for lists.
• Never promise guaranteed returns; keep compliance in mind.
• Return ONLY the requested copy – no meta commentary, no code fences.

{country_rules}

At the very end of the piece, append this italic line (no quotes):
*Past performance is not a reliable indicator of future results.*
""").strip()

# --- Trait exemplars (3 each) --------------------------------
TRAIT_EXAMPLES = {
    "Urgency": [
        "This isn't a drill — once midnight hits, your chance to secure these savings is gone forever.",
        "Time’s ticking — when the clock hits zero tonight, you’re out of luck.",
        "You have exactly one shot. Miss today’s deadline, and it's gone forever."
    ],
    "Data_Richness": [
        "Last year alone, our recommendations averaged returns 220% higher than the market average.",
        "Our analysis has identified 73% higher returns than the average ASX investor over three consecutive years.",
        "More than 85% of our recommended stocks outperformed the market last fiscal year alone."
    ],
    "Social_Proof": [
        "Thousands of investors trust Motley Fool every year to transform their financial future.",
        "Australia’s leading financial experts have rated us #1 three years in a row.",
        "Join over 125,000 smart investors who rely on Motley Fool’s stock advice every month."
    ],
    "Comparative_Framing": [
        "Think back to those who seized early opportunities in the smartphone revolution.",
        "Imagine being among the first to see Netflix’s potential in 2002. That’s the kind of opportunity we’re talking about.",
        "Just like the early days of Tesla, these stocks could define your investing success for years."
    ],
    "Imagery": [
        "When that switch flips, the next phase could accelerate even faster.",
        "Think of it as a snowball rolling downhill—small at first, but soon unstoppable.",
        "Like a rocket on the launch pad, the countdown has begun and liftoff is imminent."
    ],
    "Conversational_Tone": [
        "Look — investing can feel complicated, but what if it didn't have to be?",
        "We get it—investing can seem overwhelming. But what if you had someone guiding you every step of the way?",
        "Here’s the truth: investing doesn’t have to be complicated. Let’s simplify this together."
    ],
    "FOMO": [
        "Opportunities like these pass quickly — and regret can last forever.",
        "Don’t be the one who has to tell their friends, ‘I missed out when I had the chance.’",
        "By tomorrow, your chance to act will be history. Don’t live with that regret."
    ],
    "Repetition": [
        "This offer is for today only. Today only means exactly that: today only.",
        "Act now. This offer expires tonight. Again, it expires tonight—no exceptions.",
        "This is a limited-time deal. Limited-time means exactly that: limited-time."
    ],
}

def trait_guide(traits: dict) -> str:
    out = []
    for i, (name, score) in enumerate(traits.items(), 1):
        shots = 3 if score >= 8 else 2 if score >= 4 else 1
        examples = " / ".join(f"“{s}”" for s in TRAIT_EXAMPLES[name][:shots])
        out.append(f"{i}. {name.replace('_',' ')} ({score}/10) — e.g. {examples}")
    return "\n".join(out)

# --- Micro demos --------------------------------------------
EMAIL_MICRO = """
### Example Email
**Subject Line:** Last chance to lock in $119 Motley Fool membership  
**Greeting:** Hi Sarah,  
**Body:** Tonight at midnight, your opportunity to save 60 % disappears. Thousands of Australians already rely on our ASX stock tips—now it’s your turn. Click before the timer hits zero and start investing smarter.  
**CTA:** Activate my membership  
**Sign‑off:** The Motley Fool Australia Team
""".strip()

SALES_MICRO = """
### Example Sales Page
## Headline  
One Day Only—Unlock the Silver Pass for $119  

### Introduction  
Imagine having two extra experts on your side every month…

### Key Benefits  
- Double the stock picks, triple the insight  
- ASX, growth & dividend coverage in one pass  
- 400,000+ Aussie investors already on board  

### Detailed Body  
Scroll down and you’ll see why the Silver Pass could be your portfolio’s inflection point. But remember—the $119 price tag vanishes at 11:59 pm tonight.  

### CTA  
**Yes! Secure My Pass Now**
""".strip()

# --- Structural skeletons -----------------------------------
EMAIL_STRUCT = """
### Subject Line
### Greeting
### Body (benefits, urgency, proofs)
### Call‑to‑Action
### Sign‑off
""".strip()

SALES_STRUCT = """
## Headline
### Introduction
### Key Benefit Paragraphs
### Detailed Body
### Call‑to‑Action
""".strip()

# ────────────────────────────────────────────────────────────
# 5.  Prompt builder
# ────────────────────────────────────────────────────────────
def build_prompt(copy_type, copy_struct, traits, brief, length_choice, original=None):
    exemplar = EMAIL_MICRO if copy_type.startswith("📧") else SALES_MICRO
    hard = []
    if traits["Urgency"] >= 8:
        hard.append("- Include a deadline phrase in headline/subject **and** CTA.")
    if traits["Social_Proof"] >= 6:
        hard.append("- Provide ≥3 credibility builders (testimonial, member count, expert quote).")
    if traits["Data_Richness"] >= 7:
        hard.append("- Cite ≥1 numeric performance figure (% return, CAGR, dollar value).")
    hard_block = "#### Hard Requirements\n" + "\n".join(hard) if hard else ""
    edit_block = f"\n\n### ORIGINAL COPY\n{original}\n### END ORIGINAL" if original else ""

    # ------ length block ------------------------------------
    min_len, max_len = LENGTH_RULES[length_choice]
    if max_len:
        length_block = f"#### Length Requirement\nWrite between **{min_len} and {max_len} words**."
    else:
        length_block = f"#### Length Requirement\nWrite **at least {min_len} words**."

    return f"""
{trait_guide(traits)}

{exemplar}

#### Structure to Follow
{copy_struct}

{hard_block}

#### Campaign Brief
{line('Hook', brief['hook'])}{line('Details', brief['details'])}{line('Offer', f"Special {brief['offer_price']} (Retail {brief['retail_price']}), Term {brief['offer_term']}")}{line('Reports', brief['reports'])}{line('Stocks to Tease', brief['stocks_to_tease'])}{line('Quotes/News', brief['quotes_news'])}

{length_block}

Please limit bullet lists to three or fewer and favour full‑sentence paragraphs elsewhere.

### END INSTRUCTIONS
""".strip()

# ────────────────────────────────────────────────────────────
# 6.  Unified LLM helper
# ────────────────────────────────────────────────────────────
def run_chat(messages, stream=False, expect_json=False, max_tokens=MAX_OUTPUT_TOKENS):
    for attempt in range(5):
        try:
            if stream:
                resp = client.chat.completions.create(model=OPENAI_MODEL,
                                                      messages=messages,
                                                      stream=True,
                                                      max_tokens=max_tokens)
                ph, text = st.empty(), ""
                for c in resp:
                    text += c.choices[0].delta.content or ""
                    ph.markdown(text)
                return text.strip()
            else:
                kwargs = {"max_tokens": max_tokens}
                if expect_json:
                    kwargs["response_format"] = {"type": "json_object"}
                resp = client.chat.completions.create(model=OPENAI_MODEL,
                                                      messages=messages,
                                                      **kwargs)
                return resp.choices[0].message.content.strip()
        except Exception:
            time.sleep(2 ** attempt)

# ────────────────────────────────────────────────────────────
# 7A.  AI Pair‑editor
# ────────────────────────────────────────────────────────────
def self_qa(draft, copy_type):
    if not AUTO_QA:
        return draft

    # ----- dynamic length enforcement -----------------------
    min_len, _ = LENGTH_RULES.get(st.session_state.length_choice, (0, None))
    if min_len and len(draft.split()) < min_len:
        crit = f"- Draft is only {len(draft.split())} words (< {min_len}). Please expand."
    else:
        crit = ""

    if not crit:  # run normal checks only when length passes
        crit = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[{"role":"system","content":"You are an obsessive editorial QA bot."},
                      {"role":"user","content":f"""
Check copy for:
• Hard requirements
• Structure matches {copy_type}
• Disclaimer present
Return ONLY “PASS” or bullet fixes.
--- COPY ---
{draft}
--- END ---
"""}]
        ).choices[0].message.content

    if "PASS" in crit.upper():
        return draft

    patched = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[{"role":"system","content":"Revise copy to address feedback."},
                  {"role":"user","content":f"""
Apply fixes, output full revised copy ONLY.
### FIXES
{crit}
### ORIGINAL
{draft}
"""}]
    ).choices[0].message.content.strip()
    return patched

# ────────────────────────────────────────────────────────────
# 7B.  Variant generator helper (unchanged)
# ────────────────────────────────────────────────────────────
def generate_variants(base_copy: str, n: int = 5):
    prompt = f"""
Write {n} alternative subject‑line/headline ideas AND {n} alternative CTA button labels
for the copy below, preserving tone and urgency.
Return JSON: {{ "headlines": [...], "ctas": [...] }}

--- COPY ---
{base_copy}
--- END COPY ---
"""
    resp = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[{"role":"system","content":"You are a world‑class copywriter."},
                  {"role":"user","content":prompt}],
        response_format={"type":"json_object"},
        temperature=0.8
    )
    return json.loads(resp.choices[0].message.content)

# ────────────────────────────────────────────────────────────
# 8.  UI – Generate tab
# ────────────────────────────────────────────────────────────
tab_gen, tab_adapt = st.tabs(["✍️ Generate Copy", "🌐 Adapt Copy"])

with tab_gen:
    # --- Sliders
    with st.sidebar.expander("🎚️ Linguistic Trait Intensity", True):
        with st.form("trait_form"):
            trait_scores = {
                "Urgency":             st.slider("Urgency & Time Sensitivity", 1, 10, 8),
                "Data_Richness":       st.slider("Data‑Richness & Numerical Emphasis", 1, 10, 7),
                "Social_Proof":        st.slider("Social Proof & Testimonials", 1, 10, 6),
                "Comparative_Framing": st.slider("Comparative Framing", 1, 10, 6),
                "Imagery":             st.slider("Imagery & Metaphors", 1, 10, 7),
                "Conversational_Tone": st.slider("Conversational Tone", 1, 10, 8),
                "FOMO":                st.slider("FOMO", 1, 10, 7),
                "Repetition":          st.slider("Repetition for Emphasis", 1, 10, 5),
            }
            update_traits = st.form_submit_button("🔄 Update Copy")

    # --- Inputs
    country   = st.selectbox("🌐 Target Country", list(COUNTRY_RULES))
    copy_type = st.selectbox("Copy Type", ["📧 Email", "📝 Sales Page"])
    length_choice = st.selectbox("Desired Length", list(LENGTH_RULES))
    st.session_state.length_choice = length_choice

    st.subheader("Campaign Brief")
    hook    = st.text_area("🪝 Campaign Hook")
    details = st.text_area("📦 Product / Offer Details")

    c1, c2, c3 = st.columns(3)
    offer_price  = c1.text_input("Special Offer Price")
    retail_price = c2.text_input("Retail Price")
    offer_term   = c3.text_input("Subscription Term")

    reports         = st.text_area("📑 Included Reports")
    stocks_to_tease = st.text_input("📈 Stocks to Tease (optional)")
    st.subheader("📰 Quotes or Recent News (optional)")
    quotes_news = st.text_area("Add quotes, stats, or timely news to reference")

    show_critique = st.checkbox("🧐 Show AI critique after draft", value=False)

    def brief():
        return {"country": country, "hook": hook, "details": details,
                "offer_price": offer_price, "retail_price": retail_price,
                "offer_term": offer_term, "reports": reports,
                "stocks_to_tease": stocks_to_tease, "quotes_news": quotes_news}

    copy_struct = EMAIL_STRUCT if copy_type.startswith("📧") else SALES_STRUCT

    # ────────── Core generator ────────── #
    def generate(old=None):
        prompt_core = build_prompt(copy_type, copy_struct,
                                   trait_scores, brief(), length_choice, old)
        user_instr = dedent("""
        ### TASK
        1. Create a concise INTERNAL bullet plan covering:
           • Hook & opening flow
           • Placement of proof, urgency, CTA
           • Any standout stats, metaphors, social proof you intend to use
        2. Then write the final copy.

        Respond ONLY as valid JSON with exactly two keys:
        {
          "plan": "<the bullet outline>",
          "copy": "<the finished marketing copy>"
        }
        """).strip()

        msgs = [
            {"role":"system",
             "content": SYSTEM_PROMPT.format(country_rules=COUNTRY_RULES[country])},
            {"role":"user",
             "content": user_instr + "\n\n" + prompt_core}
        ]

        with st.spinner("Crafting copy…"):
            raw_json = run_chat(msgs, expect_json=True)

        try:
            data = json.loads(raw_json)
        except json.JSONDecodeError:
            data = {"plan": "", "copy": raw_json}

        st.session_state.internal_plan = data["plan"].strip()
        draft = self_qa(data["copy"].strip(), copy_type)

        if show_critique:
            crit = client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[
                    {"role":"system","content":"Give concise, constructive feedback."},
                    {"role":"user","content":f"""
In 3 bullets – one strength, one weakness, one improvement.
--- COPY ---
{draft}
--- END ---
"""}]
            ).choices[0].message.content
            st.info(crit)

        return draft

    # --- Buttons
    if st.button("✨ Generate Copy"):
        st.session_state.generated_copy = generate()

    if update_traits and st.session_state.generated_copy:
        st.session_state.generated_copy = generate(st.session_state.generated_copy)

    # --- Display
    if st.session_state.generated_copy:
        st.subheader("📝 Current Copy")
        st.markdown(st.session_state.generated_copy)

        # with st.expander("🔍 Internal Plan"):
        #     st.markdown(st.session_state.internal_plan)

        st.code(st.session_state.generated_copy, language="markdown")

        # variant grid
        if st.button("🎯 Generate 5 Alt Headlines & CTAs"):
            with st.spinner("Brainstorming variants…"):
                variants = generate_variants(st.session_state.generated_copy)

            st.subheader("📰 Headline Ideas")
            cols = st.columns(5)
            for i, text in enumerate(variants["headlines"]):
                with cols[i]:
                    st.markdown(f"**{i+1}.** {text}")
                    st.radio(f"head_{i}", ["👍", "👎"], horizontal=True, label_visibility="collapsed")

            st.subheader("🔘 CTA Button Ideas")
            cols = st.columns(5)
            for i, text in enumerate(variants["ctas"]):
                with cols[i]:
                    st.markdown(f"**{i+1}.** {text}")
                    st.radio(f"cta_{i}", ["👍", "👎"], horizontal=True, label_visibility="collapsed")

        col1, col2 = st.columns(2)
        if col1.button("💾 Save DOCX", key="gen_save"):
            doc = Document(); doc.add_paragraph(st.session_state.generated_copy)
            buf = BytesIO(); doc.save(buf); buf.seek(0)
            st.download_button("📥 Download DOCX", buf, "mf_copy.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        if col2.button("🗑️ Clear"):
            st.session_state.generated_copy = ""
            st.session_state.internal_plan = ""
            st.experimental_rerun()

# ────────────────────────────────────────────────────────────
# 9.  UI – Adapt tab (unchanged)
# ────────────────────────────────────────────────────────────
with tab_adapt:
    st.markdown("### Paste the original copy and select a **target country**.")
    original_text = st.text_area("Original Copy", height=250)

    colA, colB = st.columns(2)
    source_c = colA.selectbox("Original Country", list(COUNTRY_RULES))
    target_c = colB.selectbox("Target Country",
                              [c for c in COUNTRY_RULES if c != source_c])

    if st.button("🌐 Adapt Copy") and original_text.strip():
        msgs = [
            {"role":"system",
             "content": SYSTEM_PROMPT.format(country_rules=COUNTRY_RULES[target_c])},
            {"role":"user",
             "content": (
                 f"Adapt the following marketing copy for a {target_c} audience.\n"
                 "Update spelling, currency, market references; preserve tone & structure.\n\n"
                 "--- ORIGINAL COPY START ---\n"
                 f"{original_text}\n"
                 "--- ORIGINAL COPY END ---\n"
                 "### END INSTRUCTIONS"
             )}
        ]
        with st.spinner("Adapting…"):
            st.session_state.adapted_copy = run_chat(msgs)

    if st.session_state.adapted_copy:
        st.subheader("🌐 Adapted Copy")
        st.markdown(st.session_state.adapted_copy)

        a1, a2 = st.columns(2)
        if a1.button("💾 Save DOCX", key="adapt_save"):
            doc = Document(); doc.add_paragraph(st.session_state.adapted_copy)
            buf = BytesIO(); doc.save(buf); buf.seek(0)
            st.download_button("📥 Download DOCX", buf, "mf_adapted.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        if a2.button("🗑️ Clear Adapted"):
            st.session_state.adapted_copy = ""
